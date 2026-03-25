import os
import json
import uuid
import re
import argparse
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Any, Optional

# 尝试导入所需库
try:
    from docx import Document
except ImportError:
    Document = None
    print("警告：python-docx 未安装，无法处理 Word 文件。请执行 pip install python-docx")

try:
    import openpyxl
except ImportError:
    openpyxl = None
    print("警告：openpyxl 未安装，无法处理 Excel 文件。请执行 pip install openpyxl")

# 尝试导入 calamine（用于处理 strict open XML 格式 Excel）
try:
    from python_calamine import CalamineWorkbook
    CALAMINE_AVAILABLE = True
except ImportError:
    CALAMINE_AVAILABLE = False
    print("警告：python-calamine 未安装，遇到 strict open XML 格式 Excel 时可能无法解析。请执行 pip install python-calamine")


class DocumentParser:
    """
    文档解析器：支持 .docx, .xlsx, .md, .txt 格式，输出统一 JSON 格式。
    包含文本清洗功能：去除多余空格、控制字符，合并内部空白。
    Word 解析：优先使用 python-docx，失败时自动回退到手动解析 word/document.xml。
    """

    def __init__(self):
        self.supported_extensions = {'.docx', '.xlsx', '.md', '.txt'}

    @staticmethod
    def _clean_text(text: str) -> str:
        """清洗文本：移除控制字符、合并空白、去除首尾空白"""
        if not isinstance(text, str):
            text = str(text)
        # 移除控制字符（保留 \t, \n 等）
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        # 合并内部连续空白为单个空格
        text = re.sub(r'\s+', ' ', text)
        # 去除首尾空白
        return text.strip()

    def _clean_paragraphs(self, paragraphs: List[str]) -> List[str]:
        """清洗段落列表"""
        return [self._clean_text(p) for p in paragraphs if self._clean_text(p)]

    def _clean_tables(self, tables: List[List[List[str]]]) -> List[List[List[str]]]:
        """清洗表格中的所有单元格文本"""
        cleaned_tables = []
        for table in tables:
            cleaned_table = []
            for row in table:
                cleaned_row = [self._clean_text(cell) for cell in row]
                if any(cell for cell in cleaned_row):
                    cleaned_table.append(cleaned_row)
            if cleaned_table:
                cleaned_tables.append(cleaned_table)
        return cleaned_tables

    def parse(self, file_path: Path, doc_id: Optional[str] = None) -> Dict[str, Any]:
        """解析单个文档并返回统一的 JSON 结构"""
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式：{ext}，支持格式：{self.supported_extensions}")

        # 根据扩展名调用对应的解析方法
        if ext == '.docx':
            paragraphs, tables = self._parse_docx(file_path)
        elif ext == '.xlsx':
            paragraphs, tables = self._parse_xlsx(file_path)
        elif ext == '.md':
            paragraphs, tables = self._parse_markdown(file_path)
        elif ext == '.txt':
            paragraphs, tables = self._parse_txt(file_path)
        else:
            raise RuntimeError("意外的文件类型")

        # 清洗段落和表格
        paragraphs = self._clean_paragraphs(paragraphs)
        tables = self._clean_tables(tables)

        # 生成原始文本（清洗后的段落用换行连接）
        raw_text = '\n'.join(paragraphs)

        if doc_id is None:
            doc_id = str(uuid.uuid4())

        return {
            "doc_id": doc_id,
            "paragraphs": paragraphs,
            "tables": tables,
            "raw_text": raw_text
        }

    # ---------- Word 解析（增强版，支持手动回退）----------
    def _parse_docx(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        """
        解析 Word 文档。
        优先使用 python-docx，如果因缺少部件失败，则回退到手动解析 word/document.xml。
        """
        if Document is None:
            raise ImportError("python-docx 未安装，无法解析 Word 文档")

        # 首先检查是否为有效的 ZIP 文件（.docx 本质是 ZIP）
        if not zipfile.is_zipfile(file_path):
            raise ValueError(f"文件 {file_path} 不是有效的 ZIP 归档，可能为旧版 .doc 文件，无法解析。")

        try:
            # 尝试用 python-docx 解析
            return self._parse_docx_with_python_docx(file_path)
        except KeyError as e:
            # 缺少部件（如 footnotes.xml），回退到手动解析
            print(f"  解析 {file_path.name} 时缺少部件 ({e})，尝试手动解析...")
            return self._parse_docx_manual(file_path)
        except Exception as e:
            # 其他未知错误
            raise RuntimeError(f"解析 Word 文档失败：{e}")

    def _parse_docx_with_python_docx(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        """使用 python-docx 解析（标准 .docx）"""
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []

        for table in doc.tables:
            table_data = []
            try:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if not cell_text:
                            # 通过段落拼接（保留换行）
                            cell_text = '\n'.join([p.text for p in cell.paragraphs if p.text.strip()])
                        row_data.append(cell_text)
                    if any(cell.strip() for cell in row_data):
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            except Exception as e:
                print(f"  警告：Word 表格解析失败，已跳过。错误：{e}")
                continue

        return paragraphs, tables

    def _parse_docx_manual(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        """
        手动解析 docx 文件：解压并读取 word/document.xml，提取段落和表格。
        用于处理 python-docx 无法读取的兼容模式文档。
        """
        paragraphs = []
        tables = []

        # WordprocessingML 命名空间
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        try:
            with zipfile.ZipFile(file_path, 'r') as zipf:
                # 读取主文档部分
                with zipf.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                # 提取所有段落（w:p）
                for p in root.findall('.//w:p', namespaces=ns):
                    # 获取段落中所有文本（w:t）
                    texts = p.findall('.//w:t', namespaces=ns)
                    p_text = ''.join(t.text for t in texts if t.text)
                    if p_text.strip():
                        paragraphs.append(p_text)

                # 提取所有表格（w:tbl）
                for tbl in root.findall('.//w:tbl', namespaces=ns):
                    table_data = []
                    for row in tbl.findall('.//w:tr', namespaces=ns):
                        row_data = []
                        for cell in row.findall('.//w:tc', namespaces=ns):
                            # 单元格内可能包含多个段落，提取所有文本并用换行连接
                            cell_texts = []
                            for p in cell.findall('.//w:p', namespaces=ns):
                                p_texts = p.findall('.//w:t', namespaces=ns)
                                p_text = ''.join(t.text for t in p_texts if t.text)
                                if p_text:
                                    cell_texts.append(p_text)
                            cell_text = '\n'.join(cell_texts) if cell_texts else ''
                            row_data.append(cell_text)
                        # 只保留至少有一个非空单元格的行
                        if any(cell.strip() for cell in row_data):
                            table_data.append(row_data)
                    if table_data:
                        tables.append(table_data)

        except KeyError:
            # 如果连 word/document.xml 都不存在，说明文件严重损坏
            raise RuntimeError(f"文件 {file_path} 中未找到 word/document.xml，可能已损坏。")
        except Exception as e:
            raise RuntimeError(f"手动解析 docx 失败：{e}")

        return paragraphs, tables

    # ---------- Excel 解析（增强版，支持 strict open XML）----------
    def _parse_xlsx(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        """
        解析 Excel 文件。
        首先尝试用 openpyxl 读取普通格式，如果发现工作表数为 0，
        且 calamine 可用，则自动切换至 calamine 引擎读取 strict open XML 格式。
        修改点：paragraphs 收集每行合并后的字符串，而非每个单元格单独成段。
        """
        paragraphs = []
        tables = []

        # 尝试用 openpyxl 读取
        if openpyxl is not None:
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                if len(wb.worksheets) > 0:  # 普通格式，有工作表
                    for sheet in wb.worksheets:
                        # 构建表格数据
                        table_data = []
                        for row in sheet.iter_rows(values_only=True):
                            # 将整行单元格转为字符串列表
                            row_str = [str(cell) if cell is not None else '' for cell in row]
                            # 收集合并后的行文本（用空格连接）
                            merged_row_text = ' '.join(cell for cell in row_str if cell.strip())
                            if merged_row_text:
                                paragraphs.append(merged_row_text)
                            # 保留完整表格数据
                            if any(cell.strip() for cell in row_str):
                                table_data.append(row_str)
                        if table_data:
                            tables.append(table_data)
                    return paragraphs, tables
                else:
                    print(f"  检测到文件可能为 strict open XML 格式，尝试使用 calamine 引擎...")
            except Exception as e:
                print(f"  openpyxl 读取失败：{e}，尝试使用 calamine 引擎...")
        else:
            print("  openpyxl 未安装，尝试使用 calamine 引擎...")

        # 如果 openpyxl 失败或无工作表，且 calamine 可用，则用 calamine 读取
        if not CALAMINE_AVAILABLE:
            raise ImportError("既无法使用 openpyxl，也未安装 python-calamine，无法解析该 Excel 文件。")

        try:
            workbook = CalamineWorkbook.from_path(str(file_path))
            sheet_names = workbook.sheet_names
            if not sheet_names:
                raise ValueError("Excel 文件中未找到任何工作表")

            for sheet_name in sheet_names:
                sheet = workbook.get_sheet_by_name(sheet_name)
                # to_python() 返回二维列表，元素为 Python 类型（str, float, int, datetime, None 等）
                data = sheet.to_python()

                # 构建表格数据
                table_data = []
                for row in data:
                    row_str = [str(cell) if cell is not None else '' for cell in row]
                    # 收集合并后的行文本（用空格连接）
                    merged_row_text = ' '.join(cell for cell in row_str if cell.strip())
                    if merged_row_text:
                        paragraphs.append(merged_row_text)
                    # 保留完整表格数据
                    if any(cell.strip() for cell in row_str):
                        table_data.append(row_str)
                if table_data:
                    tables.append(table_data)
        except Exception as e:
            raise RuntimeError(f"calamine 解析失败：{e}")

        return paragraphs, tables

    # ---------- Markdown 解析 ----------
    def _parse_markdown(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        content = self._read_text_file(file_path)
        lines = content.splitlines()
        paragraphs = []
        tables = []
        i = 0
        while i < len(lines):
            line = lines[i].rstrip('\n')
            if line.strip().startswith('|'):
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row = [cell.strip() for cell in lines[i].strip('| \n').split('|')]
                    table_rows.append(row)
                    i += 1
                if table_rows:
                    tables.append(table_rows)
                continue
            elif line.strip().startswith('#'):
                heading = line.lstrip('#').strip()
                if heading:
                    paragraphs.append(heading)
            elif line.strip():
                paragraphs.append(line.strip())
            i += 1
        return paragraphs, tables

    # ---------- TXT 解析 ----------
    def _parse_txt(self, file_path: Path) -> tuple[List[str], List[List[List[str]]]]:
        content = self._read_text_file(file_path)
        lines = content.splitlines()
        paragraphs = [line.strip() for line in lines if line.strip()]
        return paragraphs, []

    # ---------- 通用文本文件读取（自动检测编码）----------
    def _read_text_file(self, file_path: Path) -> str:
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(f"无法用以下编码解码文件 {file_path}：{encodings}")


def batch_parse(input_dir: Path, output_dir: Path = None, parser: DocumentParser = None):
    """
    批量解析指定目录及其所有子目录下的支持文档，并将每个文档的解析结果保存为 JSON 文件。
    输出文件名格式：[子目录路径_]原文件名_扩展名.json（子目录路径中的分隔符替换为下划线）
    同时文档 ID（doc_id）也采用相同格式（不含 .json 后缀）。
    """
    input_dir = Path(input_dir).resolve()
    if not input_dir.is_dir():
        raise NotADirectoryError(f"输入目录不存在：{input_dir}")

    if output_dir is None:
        output_dir = input_dir / "output_json"
    else:
        output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    if parser is None:
        parser = DocumentParser()

    success_count = 0
    fail_count = 0
    # 递归遍历所有子目录
    for file_path in input_dir.rglob('*'):
        if not file_path.is_file():
            continue
        ext = file_path.suffix.lower()
        if ext not in parser.supported_extensions:
            continue

        rel_path = file_path.relative_to(input_dir)
        print(f"正在解析：{rel_path}")
        try:
            # 生成包含目录信息的文档 ID（格式同输出文件名，无 .json）
            parts = rel_path.parts
            stem = file_path.stem
            dir_part = "_".join(parts[:-1]) if len(parts) > 1 else ""
            ext_part = ext[1:]  # 去掉开头的点，如 'docx'
            if dir_part:
                doc_id = f"{dir_part}_{stem}_{ext_part}"
            else:
                doc_id = f"{stem}_{ext_part}"

            result = parser.parse(file_path, doc_id=doc_id)

            # 生成输出文件名（与 doc_id 一致，但加上 .json）
            out_filename = f"{doc_id}.json"
            output_file = output_dir / out_filename
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)

            print(f"  成功，已保存至：{output_file}")
            success_count += 1
        except Exception as e:
            print(f"  失败：{e}")
            fail_count += 1

    print(f"\n批量解析完成。成功：{success_count} 个，失败：{fail_count} 个。")


def main():
    parser = argparse.ArgumentParser(description="文档解析工具：递归遍历目录，将多种格式文档转换为统一 JSON 结构（含文本清洗）")
    parser.add_argument("input_dir", type=str, help="包含待解析文档的输入目录（会递归遍历所有子目录）")
    parser.add_argument("-o", "--output_dir", type=str, default=None, help="输出目录（默认输入目录下的 output_json）")
    args = parser.parse_args()

    batch_parse(args.input_dir, args.output_dir)


if __name__ == "__main__":
    main()